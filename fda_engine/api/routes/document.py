"""Document tree and generation routes."""
from __future__ import annotations

import asyncio
import uuid

from fastapi import APIRouter, BackgroundTasks, HTTPException

from fda_engine.api.deps import get_engine, get_state
from fda_engine.api.models import (
    DocumentGenerateRequest,
    DocumentGenerateResponse,
    DocumentTreeResponse,
    DocumentTreeNode,
)

router = APIRouter(prefix="/document", tags=["document"])


def _build_510k_tree() -> DocumentTreeNode:
    """Build the 510(k) document tree structure."""
    return DocumentTreeNode(
        node_id="root",
        title="510(k) Submission Package",
        required=True,
        description="Complete FDA 510(k) premarket notification submission",
        children=[
            DocumentTreeNode(
                node_id="cover_letter",
                title="Cover Letter",
                required=True,
                description="Submission cover letter identifying the device and predicate",
            ),
            DocumentTreeNode(
                node_id="510k_summary",
                title="510(k) Summary or 510(k) Statement",
                required=True,
                description="Summary of the submitter's safety and effectiveness data",
            ),
            DocumentTreeNode(
                node_id="indication_for_use",
                title="Indications for Use Statement",
                required=True,
                description="FDA Form 3881 - Intended use of the device",
            ),
            DocumentTreeNode(
                node_id="device_description",
                title="Device Description",
                required=True,
                description="Detailed description of the device, components, and specifications",
                children=[
                    DocumentTreeNode(
                        node_id="device_specs",
                        title="Technical Specifications",
                        required=True,
                        description="Dimensions, materials, weight, and technical details",
                    ),
                    DocumentTreeNode(
                        node_id="device_diagrams",
                        title="Device Diagrams and Photos",
                        required=True,
                        description="Drawings, photos, and illustrations of the device",
                    ),
                ],
            ),
            DocumentTreeNode(
                node_id="substantial_equivalence",
                title="Substantial Equivalence Comparison",
                required=True,
                description="Comparison with predicate device(s)",
                children=[
                    DocumentTreeNode(
                        node_id="se_comparison_table",
                        title="Comparison Table",
                        required=True,
                        description="Side-by-side comparison with predicate device",
                    ),
                    DocumentTreeNode(
                        node_id="se_analysis",
                        title="Equivalence Analysis",
                        required=True,
                        description="Analysis demonstrating substantial equivalence",
                    ),
                ],
            ),
            DocumentTreeNode(
                node_id="performance_testing",
                title="Performance Testing",
                required=True,
                description="Test results demonstrating safety and effectiveness",
                children=[
                    DocumentTreeNode(
                        node_id="biocompatibility",
                        title="Biocompatibility Testing (ISO 10993)",
                        required=True,
                        description="Biocompatibility evaluation per ISO 10993",
                    ),
                    DocumentTreeNode(
                        node_id="electrical_safety",
                        title="Electrical Safety Testing (IEC 60601)",
                        required=False,
                        description="Electrical safety and EMC testing if applicable",
                    ),
                    DocumentTreeNode(
                        node_id="sterilization_validation",
                        title="Sterilization Validation",
                        required=False,
                        description="Sterilization method validation (EO, Gamma, etc.)",
                    ),
                    DocumentTreeNode(
                        node_id="shelf_life",
                        title="Shelf Life / Packaging Testing",
                        required=False,
                        description="Accelerated and real-time aging studies",
                    ),
                    DocumentTreeNode(
                        node_id="functional_testing",
                        title="Functional Performance Testing",
                        required=True,
                        description="Testing per device-specific standards",
                    ),
                ],
            ),
            DocumentTreeNode(
                node_id="biocompatibility_report",
                title="Biocompatibility Report",
                required=True,
                description="Comprehensive biocompatibility evaluation report",
            ),
            DocumentTreeNode(
                node_id="software_level",
                title="Software Documentation (Level of Concern)",
                required=False,
                description="Software description and risk classification if device contains software",
            ),
            DocumentTreeNode(
                node_id="cybersecurity",
                title="Cybersecurity Documentation",
                required=False,
                description="Cybersecurity risk management for network-connected devices",
            ),
            DocumentTreeNode(
                node_id="labeling",
                title="Labeling",
                required=True,
                description="Device labeling including IFU, labels, and packaging",
                children=[
                    DocumentTreeNode(
                        node_id="ifu",
                        title="Instructions for Use (IFU)",
                        required=True,
                        description="User instructions and safety information",
                    ),
                    DocumentTreeNode(
                        node_id="device_labels",
                        title="Device Labels",
                        required=True,
                        description="Principal display panel and additional labels",
                    ),
                ],
            ),
            DocumentTreeNode(
                node_id="sterilization",
                title="Sterilization Documentation",
                required=False,
                description="Sterilization validation and method description",
            ),
            DocumentTreeNode(
                node_id="bioburden",
                title="Bioburden Testing",
                required=False,
                description="Bioburden and sterility testing results",
            ),
            DocumentTreeNode(
                node_id="risk_management",
                title="Risk Management File (ISO 14971)",
                required=True,
                description="Risk analysis, evaluation, and control measures",
            ),
            DocumentTreeNode(
                node_id="clinical_data",
                title="Clinical Data",
                required=False,
                description="Clinical studies or literature review supporting safety and effectiveness",
            ),
            DocumentTreeNode(
                node_id="truth_of_origin",
                title="Truth of Origin Letter",
                required=True,
                description="Letter confirming the submitter's right to reference third-party data",
            ),
            DocumentTreeNode(
                node_id="declarations",
                title="Declarations and Statements",
                required=True,
                description="Truthful and accurate statement, class III certification",
            ),
        ],
    )


def _build_nmpa_tree(product_chars: dict | None = None) -> DocumentTreeNode:
    """Build the NMPA registration document tree structure.
    
    Dynamically adjusts based on product characteristics:
    - 有源/无源: determines electrical safety requirements
    - 无菌/非无菌: determines sterilization documentation
    - 软件: determines software documentation
    - 皮肤接触: determines biocompatibility scope
    """
    chars = product_chars or {}
    product_type = chars.get("product_type", "有源")
    sterility = chars.get("sterility", "非无菌")
    has_software = chars.get("has_software", False)
    has_battery = chars.get("has_battery", False)
    has_wireless = chars.get("has_wireless", False)
    # Determine if device has skin contact (for blood pressure monitors, wearable devices, etc.)
    has_skin_contact = chars.get("has_skin_contact", True)  # Default True for wearable devices
    contact_tissue = chars.get("contact_tissue", "表面")
    
    # Base documents always required
    children = [
        DocumentTreeNode(
            node_id="cover_letter",
            title="1. 申请表",
            required=True,
            description="医疗器械注册申请表，包含产品基本信息、申请人信息",
        ),
        DocumentTreeNode(
            node_id="product_info",
            title="2. 产品技术要求",
            required=True,
            description="产品技术要求，包含产品性能指标、检验方法",
            children=[
                DocumentTreeNode(
                    node_id="performance_index",
                    title="2.1 性能指标",
                    required=True,
                    description="产品各项性能指标要求",
                ),
                DocumentTreeNode(
                    node_id="test_methods",
                    title="2.2 检验方法",
                    required=True,
                    description="各项性能指标的检验方法",
                ),
            ],
        ),
        DocumentTreeNode(
            node_id="risk_management",
            title="3. 风险管理报告",
            required=True,
            description="依据YY/T 0316的风险管理报告",
            children=[
                DocumentTreeNode(
                    node_id="risk_analysis",
                    title="3.1 风险分析",
                    required=True,
                    description="产品危害识别和风险分析",
                ),
                DocumentTreeNode(
                    node_id="risk_evaluation",
                    title="3.2 风险评价",
                    required=True,
                    description="风险可接受性评价",
                ),
                DocumentTreeNode(
                    node_id="risk_control",
                    title="3.3 风险控制",
                    required=True,
                    description="风险控制措施及验证",
                ),
            ],
        ),
        DocumentTreeNode(
            node_id="safety_basic",
            title="4. 医疗器械安全和性能基本原则清单",
            required=True,
            description="符合《医疗器械安全和性能基本原则》的符合性说明",
        ),
        DocumentTreeNode(
            node_id="product_desc",
            title="5. 产品说明书和标签样稿",
            required=True,
            description="产品使用说明书、标签、包装标识样稿",
            children=[
                DocumentTreeNode(
                    node_id="ifu",
                    title="5.1 使用说明书",
                    required=True,
                    description="产品使用说明书（IFU）",
                ),
                DocumentTreeNode(
                    node_id="labels",
                    title="5.2 标签和包装标识",
                    required=True,
                    description="产品标签、包装标识设计稿",
                ),
            ],
        ),
    ]
    
    # Research data section - DYNAMICALLY generated based on product characteristics
    research_children = []
    
    # For blood pressure monitors and similar measurement devices:
    # 根据《血压计注册审查指导原则》和产品实际特征生成目录
    
    # 1. Performance testing - core requirement for measurement devices
    research_children.append(DocumentTreeNode(
        node_id="performance_testing",
        title="6.1 产品性能研究",
        required=True,
        description="产品功能性、准确性、可靠性性能研究（依据产品技术要求）",
    ))
    
    # 2. Electrical safety - for active devices with battery
    if product_type == "有源" or has_battery:
        research_children.append(DocumentTreeNode(
            node_id="electrical_safety",
            title="6.2 电气安全研究",
            required=True,
            description="依据GB 9706.1-2020《医用电气设备 第1部分：基本安全和基本性能的通用要求》",
        ))
    
    # 3. EMC - for active electronic devices
    if product_type == "有源":
        research_children.append(DocumentTreeNode(
            node_id="emc_testing",
            title="6.3 电磁兼容性研究",
            required=True,
            description="依据YY 9706.102-2021《医用电气设备 电磁兼容 要求和试验》",
        ))
    
    # 4. Software - if device has software (blood pressure monitors typically do)
    if has_software:
        research_children.append(DocumentTreeNode(
            node_id="software_lifecycle",
            title="6.4 软件研究资料",
            required=True,
            description="依据《医疗器械软件注册审查指导原则（2022年修订版）》，包含软件描述、开发过程、风险管理",
        ))
    
    # 5. Usability - for devices with patient/user interaction
    # 血压计需要用户自行操作，可用性是必须的
    research_children.append(DocumentTreeNode(
        node_id="usability",
        title="6.5 可用性研究",
        required=True,
        description="依据YY/T 1474-2016《医疗器械 可用性工程对医疗器械的应用》，包含使用场景、用户界面评估",
    ))
    
    # 6. Biocompatibility - ONLY for devices with tissue contact
    # 血压计袖带接触皮肤，需要皮肤刺激和致敏测试
    # 但不是所有生物相容性测试都适用
    if has_skin_contact:
        research_children.append(DocumentTreeNode(
            node_id="biocompatibility",
            title="6.6 生物相容性研究",
            required=True,
            description="依据GB/T 16886.1-2011，仅针对皮肤接触材料进行细胞毒性和致敏试验",
        ))
    
    # 7. Shelf life / aging - for products with expiration
    research_children.append(DocumentTreeNode(
        node_id="shelf_life",
        title="6.7 使用寿命研究",
        required=True,
        description="依据ASTM F1980加速老化试验，验证产品有效期",
    ))
    
    # 8. Packaging validation
    research_children.append(DocumentTreeNode(
        node_id="packaging",
        title="6.8 包装验证研究",
        required=True,
        description="包装完整性、运输模拟试验",
    ))
    
    # 9. Cleaning/disinfection - for reusable non-sterile devices
    if sterility == "非无菌":
        research_children.append(DocumentTreeNode(
            node_id="cleaning",
            title="6.9 清洗消毒研究",
            required=True,
            description="用户清洁消毒方法验证",
        ))
    
    # 10. Wireless performance - if device has wireless
    if has_wireless:
        research_children.append(DocumentTreeNode(
            node_id="wireless_performance",
            title="6.10 无线性能研究",
            required=True,
            description="蓝牙/WiFi通信性能、数据传输可靠性",
        ))
    
    # 11. Battery safety - if device has battery
    if has_battery:
        research_children.append(DocumentTreeNode(
            node_id="battery_safety",
            title="6.11 电池安全研究",
            required=True,
            description="锂电池安全性能、充放电安全",
        ))
    
    # 12. Clinical evaluation
    research_children.append(DocumentTreeNode(
        node_id="clinical_evaluation",
        title="6.12 临床评价资料",
        required=True,
        description="同品种医疗器械临床评价或临床试验（如需要）",
    ))
    
    children.append(DocumentTreeNode(
        node_id="research_data",
        title="6. 研究资料",
        required=True,
        description="产品研究性资料（根据产品特征动态生成）",
        children=research_children,
    ))
    
    # Testing report
    children.append(DocumentTreeNode(
        node_id="testing_reports",
        title="7. 注册检验报告",
        required=True,
        description="具有资质的检验机构出具的产品检验报告",
    ))
    
    # Software documentation - if device has software
    if has_software:
        children.append(DocumentTreeNode(
            node_id="software_doc",
            title="8. 软件研究资料",
            required=True,
            description="医疗器械软件描述文档（依据《医疗器械软件注册审查指导原则》）",
            children=[
                DocumentTreeNode(
                    node_id="software_desc",
                    title="8.1 软件描述文档",
                    required=True,
                    description="软件概述、用途、运行环境、技术架构",
                ),
                DocumentTreeNode(
                    node_id="software_dev",
                    title="8.2 软件开发过程",
                    required=True,
                    description="开发环境、开发工具、版本管理",
                ),
                DocumentTreeNode(
                    node_id="software_risk",
                    title="8.3 软件风险管理",
                    required=True,
                    description="软件风险分析及控制措施",
                ),
            ],
        ))
    
    # Cybersecurity - if device has wireless/network
    if has_wireless or has_software:
        children.append(DocumentTreeNode(
            node_id="cybersecurity",
            title="9. 网络安全文档",
            required=True,
            description="网络安全风险分析及管理（依据《医疗器械网络安全注册审查指导原则》）",
        ))
    
    # Sterilization documentation - for sterile devices
    if sterility == "无菌":
        children.append(DocumentTreeNode(
            node_id="sterilization",
            title="10. 灭菌验证资料",
            required=True,
            description="灭菌工艺验证报告（EO/辐照等）",
            children=[
                DocumentTreeNode(
                    node_id="sterilization_validation",
                    title="10.1 灭菌工艺验证",
                    required=True,
                    description="灭菌方法选择依据及验证报告",
                ),
                DocumentTreeNode(
                    node_id="bioburden",
                    title="10.2 初始污染菌检验",
                    required=True,
                    description="初始污染菌及无菌检验结果",
                ),
            ],
        ))
    
    # Other required documents
    children.append(DocumentTreeNode(
        node_id="other_docs",
        title="10. 其他资料",
        required=False,
        description="其他支持性文件",
        children=[
            DocumentTreeNode(
                node_id="similar_device",
                title="10.1 同品种医疗器械对比表",
                required=True,
                description="与已上市同品种医疗器械的对比分析",
            ),
            DocumentTreeNode(
                node_id="change_history",
                title="10.2 产品变化情况说明",
                required=False,
                description="产品设计变更历史（如适用）",
            ),
        ],
    ))
    
    # Quality management system
    children.append(DocumentTreeNode(
        node_id="quality_mgmt",
        title="11. 质量管理体系文件",
        required=True,
        description="质量管理体系核查申请表及相关文件",
    ))
    
    return DocumentTreeNode(
        node_id="root",
        title="医疗器械注册申报资料",
        required=True,
        description=f"{'二类' if chars.get('device_class') == '二类' else '三类'}{product_type}医疗器械（{sterility}）注册申报完整资料包",
        children=children,
    )


def _build_qms_tree(product_chars: dict | None = None) -> DocumentTreeNode:
    """Build the QMS (Quality Management System) document tree structure.
    
    This generates ISO 13485 quality management system documents.
    """
    chars = product_chars or {}
    
    children = [
        DocumentTreeNode(
            node_id="qms_01_quality_manual",
            title="一级文件：质量手册",
            required=True,
            description="ISO 13485:2016质量手册，包含质量方针、目标、组织架构、过程描述",
        ),
        DocumentTreeNode(
            node_id="qms_02_procedures",
            title="二级文件：程序文件",
            required=True,
            description="ISO 13485:2016要求的程序文件",
            children=[
                DocumentTreeNode(node_id="qms_02_QP001", title="QP-001 文件控制程序", required=True),
                DocumentTreeNode(node_id="qms_02_QP002", title="QP-002 记录控制程序", required=True),
                DocumentTreeNode(node_id="qms_02_QP003", title="QP-003 管理评审程序", required=True),
                DocumentTreeNode(node_id="qms_02_QP004", title="QP-004 人力资源管理程序", required=True),
                DocumentTreeNode(node_id="qms_02_QP005", title="QP-005 基础设施管理程序", required=True),
                DocumentTreeNode(node_id="qms_02_QP006", title="QP-006 设计和开发控制程序", required=True),
                DocumentTreeNode(node_id="qms_02_QP007", title="QP-007 采购控制程序", required=True),
                DocumentTreeNode(node_id="qms_02_QP008", title="QP-008 生产控制程序", required=True),
                DocumentTreeNode(node_id="qms_02_QP009", title="QP-009 监视测量设备控制程序", required=True),
                DocumentTreeNode(node_id="qms_02_QP010", title="QP-010 内部审核程序", required=True),
                DocumentTreeNode(node_id="qms_02_QP011", title="QP-011 不合格品控制程序", required=True),
                DocumentTreeNode(node_id="qms_02_QP012", title="QP-012 纠正措施程序", required=True),
                DocumentTreeNode(node_id="qms_02_QP013", title="QP-013 预防措施程序", required=True),
                DocumentTreeNode(node_id="qms_02_QP014", title="QP-014 风险管理程序", required=True),
                DocumentTreeNode(node_id="qms_02_QP015", title="QP-015 可追溯性管理程序", required=True),
                DocumentTreeNode(node_id="qms_02_QP016", title="QP-016 顾客反馈处理程序", required=True),
                DocumentTreeNode(node_id="qms_02_QP017", title="QP-017 产品召回程序", required=True),
                DocumentTreeNode(node_id="qms_02_QP018", title="QP-018 软件生命周期管理程序", required=True),
            ],
        ),
        DocumentTreeNode(
            node_id="qms_03_work_instructions",
            title="三级文件：作业指导书",
            required=True,
            description="各类别操作作业指导书",
            children=[
                DocumentTreeNode(node_id="qms_03_design", title="设计开发作业指导书", required=True),
                DocumentTreeNode(node_id="qms_03_procurement", title="采购作业指导书", required=True),
                DocumentTreeNode(node_id="qms_03_production", title="生产过程控制作业指导书", required=True),
                DocumentTreeNode(node_id="qms_03_inspection", title="检验作业指导书", required=True),
                DocumentTreeNode(node_id="qms_03_warehouse", title="仓储管理作业指导书", required=True),
                DocumentTreeNode(node_id="qms_03_delivery", title="交付与服务作业指导书", required=True),
            ],
        ),
        DocumentTreeNode(
            node_id="qms_04_quality_records",
            title="质量记录表格",
            required=True,
            description="质量管理体系所需的各种记录表格模板",
        ),
    ]
    
    return DocumentTreeNode(
        node_id="root",
        title="ISO 13485 质量管理体系文件",
        required=True,
        description="基于ISO 13485:2016的医疗器械质量管理体系文件",
        children=children,
    )


async def _generate_node_bg(node_id: str, task_id: str):
    """Background task for generating a single document node."""
    from loguru import logger

    state = get_state()
    engine = state.engine
    template = state.config.fda_template
    lang = state.config.language
    product_chars = state.product_characteristics

    try:
        if template == "nmpa":
            tree = _build_nmpa_tree(product_chars)
        elif template == "qtms_nmpa":
            tree = _build_qms_tree(product_chars)
        elif template == "qtms_fda":
            tree = _build_qms_tree(product_chars)
        else:
            tree = _build_510k_tree()

        if node_id == "all":
            nodes = _collect_all_nodes(tree)
        else:
            node = _find_node(tree, node_id)
            if not node:
                logger.error(f"Node not found: {node_id}")
                return
            nodes = [node]

        # Load regulations and standards
        import json as json_mod
        from pathlib import Path
        
        reg_file = Path(__file__).parent.parent.parent.parent / "regulations" / "fda_510k.json"
        if template in ("nmpa", "qtms_nmpa"):
            reg_file = Path(__file__).parent.parent.parent.parent / "regulations" / "nmpa_registration.json"
        
        regulations = {}
        if reg_file.exists():
            with open(reg_file, 'r', encoding='utf-8') as f:
                regulations = json_mod.load(f)
        
        std_file = Path(__file__).parent.parent.parent.parent / "regulations" / "standards.json"
        standards = {}
        if std_file.exists():
            with open(std_file, 'r', encoding='utf-8') as f:
                standards = json_mod.load(f)

        total = len(nodes)
        for i, node in enumerate(nodes):
            logger.info(f"Generating: {node.node_id} ({i+1}/{total})")
            
            # Update progress
            state.generation_status = "generating"
            state.generation_progress = {
                "current": i + 1,
                "total": total,
                "current_node": node.title,
                "completed_nodes": list(state.generated_docs.keys()),
                "message": f"正在生成: {node.title} ({i+1}/{total})"
            }

            # Check if this is a QMS document
            is_qms = node.node_id.startswith("qms_")
            
            # Get relevant standards for this node
            relevant_standards = _get_relevant_standards(node.node_id, standards, product_chars)

            # Get company products info
            company_products = getattr(state, 'company_products', [])
            product_context = ""
            if company_products:
                product_context = "\n\n【公司产品】\n" + "\n".join([f"- {p.get('name', '')}" for p in company_products])

            if is_qms:
                # QMS document generation - with both ISO 13485 and Chinese regulations
                system_prompt = (
                    f"你是一名医疗器械质量管理体系专家。\n"
                    f"请生成'{node.title}'。\n"
                    f"{node.description}\n\n"
                    "【必须遵循的标准和法规】\n"
                    "1. ISO 13485:2016《医疗器械 质量管理体系 用于法规的要求》\n"
                    "2. 《医疗器械监督管理条例》（国务院令第739号）\n"
                    "3. 《医疗器械生产监督管理办法》（国家市场监督管理总局令第53号）\n"
                    "4. 《医疗器械生产质量管理规范》及其附录\n"
                    "5. 《医疗器械注册与备案管理办法》（国家市场监督管理总局令第47号）\n"
                    "6. 《医疗器械说明书和标签管理规定》（国家食品药品监督管理总局令第6号）\n"
                    "7. YY/T 0287-2017（等同ISO 13485:2016）\n"
                    "8. YY/T 0316-2016（等同ISO 14971:2019）\n\n"
                    "【格式要求】\n"
                    "- 使用正式的质量管理体系文件格式\n"
                    "- 同时引用ISO 13485条款号和中国法规条款号\n"
                    "- 包含详细的操作步骤和职责说明\n"
                    "- 适用于公司所有医疗器械产品\n"
                    "- 使用Markdown格式，包含标题、列表、表格\n"
                    "- 文件编号格式：XX-QP-XXX（公司简称-程序文件-序号）\n\n"
                    "【占位符规范】\n"
                    "使用 [需填写：xxx] 标记需要用户填写的内容\n"
                )
                user_msg = f"请生成'{node.title}'。{product_context}"
            elif lang == "zh" or template in ("nmpa", "qtms_nmpa"):
                system_prompt = (
                    f"你是一名医疗器械注册申报资料撰写专家，正在生成'{node.title}'部分。\n"
                    f"{node.description}\n\n"
                    "请使用正式、专业的中文医疗器械注册申报语言撰写。\n\n"
                    "【必须遵循的法规和标准】\n"
                    "1. 《医疗器械注册与备案管理办法》（国家市场监督管理总局令第47号）\n"
                    "2. 《医疗器械安全和性能基本原则》\n"
                    "3. 《医疗器械说明书和标签管理规定》\n"
                    "4. 《医疗器械临床评价技术指导原则》\n"
                )
                if relevant_standards:
                    system_prompt += "5. 相关标准：\n"
                    for std in relevant_standards[:5]:
                        system_prompt += f"   - {std}\n"
                system_prompt += (
                    "\n【格式要求】\n"
                    "- 使用正式的法规文件格式\n"
                    "- 包含具体的标准条款引用\n"
                    "- 数据表格使用Markdown格式，用 | 分隔列\n"
                    "- 使用 ## 标题分级\n"
                    "- 使用 **加粗** 突出关键信息\n"
                    "- 使用 1. 2. 3. 有序列表\n"
                    "- 使用 - 无序列表\n"
                    "- 段落之间空一行\n"
                    "- 必须包含日期和版本号占位符\n\n"
                    "【占位符规范】\n"
                    "对于缺少的具体信息，请使用以下占位符标记：\n"
                    "- [需填写：xxx] - 需要用户填写的具体内容\n"
                    "- [待补充：xxx] - 需要补充的资料\n"
                    "- [请插入：xxx] - 需要插入的图片或附件\n"
                    "- [日期待定] - 需要填写的日期\n"
                    "- [签名待签] - 需要签名的位置\n"
                    "例如：[需填写：产品注册证编号]、[待补充：临床试验报告摘要]、[请插入：产品铭牌照片]\n"
                )
                user_msg = f"请生成'{node.title}'部分。"
            else:
                system_prompt = (
                    f"You are a FDA regulatory writer generating the '{node.title}' section "
                    f"for a 510(k) submission. {node.description}\n\n"
                    "MANDATORY REGULATORY REFERENCES:\n"
                    "1. 21 CFR Part 807 - Premarket Notification\n"
                    "2. 21 CFR Part 820 - Quality System Regulation\n"
                    "3. FDA Guidance: Format and Content of a 510(k) Submission\n"
                    "4. FDA Guidance: Factors to Consider for Class II Special Controls\n"
                )
                if relevant_standards:
                    system_prompt += "5. Applicable Standards:\n"
                    for std in relevant_standards[:5]:
                        system_prompt += f"   - {std}\n"
                system_prompt += (
                    "\nFORMAT REQUIREMENTS:\n"
                    "- Use formal regulatory document format\n"
                    "- Include specific standard clause references\n"
                    "- Use Markdown tables for comparison data (with | separators)\n"
                    "- Use ## for headings\n"
                    "- Use **bold** for emphasis\n"
                    "- Use 1. 2. 3. numbered lists\n"
                    "- Use - bullet lists\n"
                    "- Add blank lines between paragraphs\n"
                    "- Include date and version placeholders\n"
                    "- Follow FDA eSTAR format when applicable\n\n"
                    "PLACEHOLDER CONVENTIONS:\n"
                    "For missing specific information, use these placeholders:\n"
                    "- [TO BE FILLED: xxx] - Specific content needed from user\n"
                    "- [TO BE ADDED: xxx] - Supporting documents needed\n"
                    "- [INSERT: xxx] - Images or attachments to insert\n"
                    "- [DATE PENDING] - Date to be filled\n"
                    "- [SIGNATURE REQUIRED] - Signature line\n"
                    "Examples: [TO BE FILLED: Device Registration Number], [TO BE ADDED: Clinical Study Summary]\n"
                )
                user_msg = f"Generate the '{node.title}' section."

            # Add product details from source documents
            product_details = getattr(state, 'product_details', {})
            if product_details and 'error' not in product_details:
                import json as json_mod
                details_str = json_mod.dumps(product_details, indent=2, ensure_ascii=False)
                if lang == "zh":
                    user_msg += f"\n\n【从产品资料中提取的实际信息】\n请基于以下真实产品信息生成文档，不要使用模板占位符：\n\n{details_str}"
                else:
                    user_msg += f"\n\n【ACTUAL PRODUCT INFORMATION FROM SOURCE DOCUMENTS】\nGenerate the document using the following real product data, do NOT use template placeholders:\n\n{details_str}"
            elif state.truth_params:
                import json
                params_json = json.dumps(state.truth_params, indent=2, ensure_ascii=False)
                if lang == "zh":
                    user_msg += f"\n\n关键产品参数:\n{params_json}"
                else:
                    user_msg += f"\n\nKey product parameters:\n{params_json}"

            content = await engine.generate_document(
                system_prompt=system_prompt,
                user_message=user_msg,
                max_iterations=state.config.workflow.max_iterations_per_node,
            )
            
            # Clean introductory text
            content = _clean_generated_content(content)

            state.generated_docs[node.node_id] = content
            node.status = "generated"

        state.generation_status = "completed"
        state.generation_progress["message"] = f"全部完成！共生成 {len(state.generated_docs)} 个文档"
        state.generation_progress["completed_nodes"] = list(state.generated_docs.keys())
        logger.info(f"Generation complete: {len(state.generated_docs)} documents")
    except Exception as e:
        state.generation_status = "error"
        state.generation_progress["message"] = f"生成失败: {str(e)}"
        logger.exception(f"Generation failed: {e}")


def _get_relevant_standards(node_id: str, standards: dict, product_chars: dict) -> list[str]:
    """Get relevant standards for a specific document node."""
    relevant = []
    
    # Mapping of node_id to standard categories
    node_standard_map = {
        "biocompatibility": ["biocompatibility"],
        "electrical_safety": ["electrical_safety"],
        "emc_testing": ["electrical_safety"],
        "software_doc": ["software"],
        "risk_management": ["risk_management"],
        "sterilization": ["sterilization"],
        "sterilization_validation": ["sterilization"],
        "shelf_life": ["packaging_shelf_life"],
        "cybersecurity": ["cybersecurity"],
        "device_description": ["usability"],
        "testing_reports": ["electrical_safety", "biocompatibility", "emc"],
        "research_data": ["biocompatibility", "electrical_safety", "clinical"],
    }
    
    # Get categories for this node
    categories = node_standard_map.get(node_id, [])
    
    # Add standards from each category
    for category in categories:
        if category in standards:
            for std in standards[category]:
                std_text = f"{std.get('standard', '')} - {std.get('title', '')}"
                if std_text not in relevant:
                    relevant.append(std_text)
    
    # Add product-specific standards
    product_type = product_chars.get("product_type", "有源")
    sterility = product_chars.get("sterility", "非无菌")
    has_software = product_chars.get("has_software", False)
    
    if product_type == "有源" and "electrical_safety" not in categories:
        # Add basic electrical safety standards for active devices
        if "electrical_safety" in standards:
            for std in standards["electrical_safety"][:1]:
                std_text = f"{std.get('standard', '')} - {std.get('title', '')}"
                if std_text not in relevant:
                    relevant.append(std_text)
    
    return relevant[:8]  # Limit to 8 standards


def _collect_all_nodes(node: DocumentTreeNode) -> list[DocumentTreeNode]:
    """Flatten tree into list of all nodes."""
    nodes = [node]
    for child in node.children:
        nodes.extend(_collect_all_nodes(child))
    return nodes


def _find_node(node: DocumentTreeNode, node_id: str) -> DocumentTreeNode | None:
    """Find a node by ID in the tree."""
    if node.node_id == node_id:
        return node
    for child in node.children:
        found = _find_node(child, node_id)
        if found:
            return found
    return None


@router.get("/tree", response_model=DocumentTreeResponse)
async def get_document_tree():
    """Get the document tree structure based on selected template and product characteristics."""
    state = get_state()
    template = state.config.fda_template
    product_chars = state.product_characteristics

    if template == "510k":
        root = _build_510k_tree()
    elif template == "nmpa":
        root = _build_nmpa_tree(product_chars)
    elif template == "qtms_nmpa":
        # QTMS (NMPA) generates QMS documents
        root = _build_qms_tree(product_chars)
    elif template == "qtms_fda":
        # QTMS (FDA) generates QMS documents
        root = _build_qms_tree(product_chars)
    else:
        raise HTTPException(status_code=400, detail=f"Template '{template}' not yet implemented")

    # Update statuses from state
    for node in _collect_all_nodes(root):
        if node.node_id in state.generated_docs:
            node.status = "generated"

    return DocumentTreeResponse(template=template, root=root)


@router.get("/progress")
async def get_generation_progress():
    """Get document generation progress."""
    state = get_state()
    return {
        "status": state.generation_status,
        "progress": state.generation_progress,
        "generated_count": len(state.generated_docs),
    }


@router.get("/regulations/{node_id}")
async def get_document_regulations(node_id: str):
    """Get applicable regulations and standards for a specific document node."""
    state = get_state()
    template = state.config.fda_template
    product_chars = state.product_characteristics
    
    import json as json_mod
    from pathlib import Path
    
    # Load standards
    std_file = Path(__file__).parent.parent.parent.parent / "regulations" / "standards.json"
    standards = {}
    if std_file.exists():
        with open(std_file, 'r', encoding='utf-8') as f:
            standards = json_mod.load(f).get("standards", {})
    
    relevant = _get_relevant_standards(node_id, standards, product_chars)
    
    # Get specific regulation requirements
    reg_requirements = {
        "510k": {
            "cover_letter": ["21 CFR 807.87", "FDA Form 3514"],
            "510k_summary": ["21 CFR 807.92"],
            "indication_for_use": ["FDA Form 3881"],
            "device_description": ["21 CFR 807.87(a)"],
            "substantial_equivalence": ["21 CFR 807.87(b)", "FDA SE Guidance"],
            "performance_testing": ["21 CFR 807.87(c)", "FDA Recognized Consensus Standards"],
            "labeling": ["21 CFR 801", "21 CFR 807.87(e)"],
            "risk_management": ["ISO 14971:2019", "FDA QMSR"],
        },
        "nmpa": {
            "cover_letter": ["《医疗器械注册与备案管理办法》第十条"],
            "product_info": ["《医疗器械产品技术要求编写指导原则》"],
            "risk_management": ["YY/T 0316-2016", "《医疗器械风险管理对医疗器械的应用》"],
            "safety_basic": ["《医疗器械安全和性能基本原则》"],
            "product_desc": ["《医疗器械说明书和标签管理规定》"],
            "biocompatibility": ["GB/T 16886.1-2011", "《医疗器械生物学评价》"],
            "electrical_safety": ["GB 9706.1-2020", "《医用电气设备安全要求》"],
            "emc_testing": ["YY 9706.102-2021"],
            "clinical_evaluation": ["《医疗器械临床评价技术指导原则》"],
            "software_doc": ["《医疗器械软件注册审查指导原则》"],
            "quality_mgmt": ["《医疗器械生产质量管理规范》"],
        }
    }
    
    template_regs = reg_requirements.get(template, {})
    node_regs = template_regs.get(node_id, [])
    
    return {
        "node_id": node_id,
        "template": template,
        "regulations": node_regs,
        "standards": relevant,
    }


@router.get("/product-characteristics")
async def get_product_characteristics():
    """Get analyzed product characteristics."""
    state = get_state()
    return {
        "characteristics": state.product_characteristics,
        "workspace_path": str(state.workspace_path) if state.workspace_path else None,
    }


@router.get("/content/{node_id}")
async def get_document_content(node_id: str):
    """Get the content of a generated document."""
    state = get_state()
    
    if node_id not in state.generated_docs:
        raise HTTPException(status_code=404, detail=f"Document '{node_id}' not found or not generated yet.")
    
    return {"node_id": node_id, "content": state.generated_docs[node_id]}


def _clean_generated_content(content: str) -> str:
    """Remove introductory text from generated content, keep only the actual document."""
    lines = content.split('\n')
    start_idx = 0
    
    # Skip lines that look like introductory text
    skip_patterns = [
        '好的，', '作为一名', '我将', '为您生成', '本手册', '通用地适用于',
        '作为示例', '确保内容', '开头部分', '首先，', '下面',
        '好的！', '我来', '以下', '根据您的', '基于您',
    ]
    
    for i, line in enumerate(lines):
        line_stripped = line.strip()
        # Skip empty lines at the start
        if not line_stripped and start_idx == i:
            start_idx = i + 1
            continue
        # Skip introductory patterns
        if any(pattern in line_stripped for pattern in skip_patterns):
            if i < 5:  # Only skip in first few lines
                start_idx = i + 1
                continue
        # If we hit a heading or content, stop skipping
        if line_stripped.startswith('#') or line_stripped.startswith('**') or (line_stripped and i > start_idx):
            break
    
    return '\n'.join(lines[start_idx:])


def _remove_ai_patterns(content: str) -> str:
    """Remove common AI-generated patterns to make text more natural."""
    import re
    
    # AI phrases to remove or replace
    ai_patterns = [
        (r'需要注意的是[，,]?', ''),
        (r'值得注意的是[，,]?', ''),
        (r'需要强调的是[，,]?', ''),
        (r'总的来说[，,]?', ''),
        (r'综上所述[，,]?', ''),
        (r'总之[，,]?', ''),
        (r'首先[，,]?\s*', ''),
        (r'其次[，,]?\s*', ''),
        (r'最后[，,]?\s*', ''),
        (r'此外[，,]?\s*', ''),
        (r'另外[，,]?\s*', ''),
        (r'同时[，,]?\s*', ''),
        (r'因此[，,]?\s*', ''),
        (r'所以[，,]?\s*', ''),
        (r'然而[，,]?\s*', ''),
        (r'不过[，,]?\s*', ''),
        (r'当然[，,]?\s*', ''),
        (r'事实上[，,]?\s*', ''),
        (r'实际上[，,]?\s*', ''),
        (r'一般来说[，,]?\s*', ''),
        (r'通常情况下[，,]?\s*', ''),
        (r'在大多数情况下[，,]?\s*', ''),
        (r'从本质上讲[，,]?\s*', ''),
        (r'从根本上说[，,]?\s*', ''),
        (r'根据相关法规要求[，,]?\s*', ''),
        (r'依据相关规定[，,]?\s*', ''),
        (r'根据实际情况[，,]?\s*', ''),
        (r'具体而言[，,]?\s*', ''),
        (r'简而言之[，,]?\s*', ''),
        (r'换句话说[，,]?\s*', ''),
        (r'这意味着[，,]?\s*', ''),
        (r'这表明[，,]?\s*', ''),
        (r'由此可见[，,]?\s*', ''),
        (r'据此[，,]?\s*', ''),
    ]
    
    for pattern, replacement in ai_patterns:
        content = re.sub(pattern, replacement, content)
    
    # Remove double spaces
    content = re.sub(r'  +', ' ', content)
    
    # Remove empty lines more than 2 consecutive
    content = re.sub(r'\n{3,}', '\n\n', content)
    
    return content


@router.post("/review")
async def review_documents():
    """Review all generated documents for quality and consistency."""
    from loguru import logger
    
    state = get_state()
    
    if not state.generated_docs:
        raise HTTPException(status_code=400, detail="没有已生成的文档")
    
    engine = state.engine
    lang = state.config.language
    template = state.config.fda_template
    
    # Collect all documents for review
    all_docs = []
    for node_id, content in state.generated_docs.items():
        if not node_id.startswith("qms_"):  # Skip QMS docs for now
            all_docs.append({"id": node_id, "content": content[:2000]})  # First 2000 chars
    
    if not all_docs:
        return {"status": "error", "message": "没有可审核的文档"}
    
    # Build review prompt
    docs_summary = "\n\n".join([f"【{d['id']}】\n{d['content']}" for d in all_docs[:5]])  # Review first 5 docs
    
    if lang == "zh":
        system_prompt = (
            "你是一名医疗器械注册申报资料审核专家。\n"
            "请审核以下生成的文档，找出问题并给出修改建议。\n\n"
            "【审核要点】\n"
            "1. 内容完整性：是否有遗漏的重要信息\n"
            "2. 法规符合性：是否符合相关法规要求\n"
            "3. 参数一致性：各文档中的产品参数是否一致\n"
            "4. 术语准确性：专业术语是否使用正确\n"
            "5. 格式规范性：是否符合申报资料格式要求\n"
            "6. AI痕迹：是否有明显的AI生成痕迹需要修改\n\n"
            "【输出格式】\n"
            "请返回JSON格式：\n"
            "{\n"
            '  "issues": [\n'
            '    {"doc_id": "文档ID", "issue": "问题描述", "severity": "高/中/低", "suggestion": "修改建议"}\n'
            "  ],\n"
            '  "ai_patterns": ["需要修改的AI痕迹列表"],\n'
            '  "consistency_check": {"参数名": "是否一致"},\n'
            '  "overall_score": 85,\n'
            '  "summary": "总体评价"\n'
            "}\n"
        )
    else:
        system_prompt = (
            "You are a FDA regulatory document review expert.\n"
            "Review the following generated documents for quality and consistency.\n\n"
            "REVIEW POINTS:\n"
            "1. Content completeness\n"
            "2. Regulatory compliance\n"
            "3. Parameter consistency across documents\n"
            "4. Terminology accuracy\n"
            "5. Format compliance\n"
            "6. AI-generated patterns to remove\n\n"
            "Return JSON with issues, AI patterns, consistency check, and overall score.\n"
        )
    
    try:
        result = await engine.generate_document(
            system_prompt=system_prompt,
            user_message=f"请审核以下文档：\n\n{docs_summary}",
            max_iterations=30,
        )
        
        import json
        try:
            review_result = json.loads(result)
        except json.JSONDecodeError:
            review_result = {
                "issues": [],
                "ai_patterns": [],
                "overall_score": 70,
                "summary": result[:500]
            }
        
        # Apply AI pattern removal to all documents
        fixed_count = 0
        for node_id in list(state.generated_docs.keys()):
            if not node_id.startswith("qms_"):
                original = state.generated_docs[node_id]
                cleaned = _remove_ai_patterns(original)
                if cleaned != original:
                    state.generated_docs[node_id] = cleaned
                    fixed_count += 1
        
        return {
            "status": "success",
            "review": review_result,
            "fixed_documents": fixed_count,
            "message": f"审核完成，已修复 {fixed_count} 个文档的AI痕迹"
        }
    except Exception as e:
        logger.exception(f"Review failed: {e}")
        return {"status": "error", "message": str(e)}


@router.post("/generate-qms")
async def generate_quality_management_system(req: DocumentGenerateRequest, bg: BackgroundTasks):
    """Generate ISO 13485 Quality Management System documents."""
    import json as json_mod
    from pathlib import Path
    
    state = get_state()
    template = state.config.fda_template
    lang = state.config.language
    product_chars = state.product_characteristics
    
    # Load ISO 13485 requirements
    qms_file = Path(__file__).parent.parent.parent.parent / "regulations" / "iso13485.json"
    qms_requirements = {}
    if qms_file.exists():
        with open(qms_file, 'r', encoding='utf-8') as f:
            qms_requirements = json_mod.load(f)
    
    # Generate QMS documents
    engine = state.engine
    
    async def _generate_qms_docs():
        from loguru import logger
        
        # Get company product information from state
        company_products = getattr(state, 'company_products', [])
        product_info = getattr(state, 'product_details', {})
        
        chapters = qms_requirements.get("iso13485", {}).get("chapters", {})
        
        # Product context for QMS generation
        product_context = ""
        if company_products:
            product_context = f"\n\n【公司产品清单】\n公司生产以下医疗器械产品：\n"
            for i, p in enumerate(company_products, 1):
                product_context += f"{i}. {p.get('name', '产品')} - {p.get('type', '类型')} - {p.get('description', '')}\n"
        elif product_info:
            product_context = f"\n\n【当前产品信息】\n产品名称：{product_info.get('product_name', '待填写')}\n产品类型：{product_info.get('product_type', '待填写')}\n"
        
        # Get comprehensive company info
        company_info = getattr(state, 'company_info', {})
        company_context = ""
        if company_info:
            company_context = f"\n\n【公司详细信息】\n"
            company_context += f"公司名称：{company_info.get('name', '待填写')}\n"
            company_context += f"公司规模：{company_info.get('size', '待填写')}\n"
            company_context += f"员工人数：{company_info.get('employees', '待填写')}\n"
            company_context += f"注册地址：{company_info.get('address', '待填写')}\n"
            company_context += f"生产地址：{company_info.get('productionAddress', '待填写')}\n"
            company_context += f"主要业务：{company_info.get('business', '待填写')}\n"
            
            depts = company_info.get('departments', {})
            if depts:
                company_context += "\n【组织架构】\n"
                if depts.get('management'): company_context += f"管理层：{depts['management']}\n"
                if depts.get('rnd'): company_context += f"研发部：{depts['rnd']}\n"
                if depts.get('production'): company_context += f"生产部：{depts['production']}\n"
                if depts.get('quality'): company_context += f"质量部：{depts['quality']}\n"
                if depts.get('procurement'): company_context += f"采购部：{depts['procurement']}\n"
                if depts.get('sales'): company_context += f"销售部：{depts['sales']}\n"
            
            products = company_info.get('products', [])
            if products:
                company_context += f"\n【公司产品清单】（共{len(products)}个产品）\n"
                for i, p in enumerate(products, 1):
                    company_context += f"{i}. {p.get('name', '')} - {p.get('type', '')} - {p.get('category', '')}\n"
        
        # Common regulations reference
        regulations_ref = (
            "\n\n【必须遵循的法规和标准】\n"
            "1. ISO 13485:2016《医疗器械 质量管理体系 用于法规的要求》\n"
            "2. 《医疗器械监督管理条例》（国务院令第739号）\n"
            "3. 《医疗器械生产监督管理办法》（国家市场监督管理总局令第53号）\n"
            "4. 《医疗器械生产质量管理规范》及其附录\n"
            "5. 《医疗器械注册与备案管理办法》（国家市场监督管理总局令第47号）\n"
            "6. YY/T 0287-2017（等同ISO 13485:2016）\n"
            "7. YY/T 0316-2016（等同ISO 14971:2019）\n"
        )
        
        # 1. Generate Quality Manual - DETAILED VERSION
        if chapters.get("quality_manual"):
            system_prompt = (
                "你是一名资深的医疗器械质量管理体系专家。\n"
                "请生成一份完整的、详细的质量手册。\n\n"
                "【重要要求】\n"
                "- 直接输出质量手册正文，不要任何前言、说明、介绍\n"
                "- 内容要详细、专业、可直接用于审核\n"
                "- 结合公司实际产品（如：有源医疗器械、含软件、含无线、非无菌）\n"
                "- 每个章节都要有详细的内容描述，不能只是标题\n\n"
                f"【必须遵循的法规和标准】{regulations_ref}\n\n"
                "【质量手册必须包含的章节及详细内容】\n\n"
                "1. 质量手册管理\n"
                "   - 版本控制、修改记录、发放控制\n\n"
                "2. 公司概况\n"
                "   - 公司简介、组织架构、地理位置、人员规模\n\n"
                "3. 质量方针和质量目标\n"
                "   - 质量方针声明（符合法规要求）\n"
                "   - 年度质量目标（可量化的指标）\n"
                "   - 目标分解到各部门\n\n"
                "4. 质量管理体系\n"
                "   - 体系范围（覆盖所有产品）\n"
                "   - 过程识别和相互作用\n"
                "   - 过程方法的应用\n"
                "   - 体系文件结构（一级、二级、三级）\n\n"
                "5. 管理职责\n"
                "   - 最高管理者承诺\n"
                "   - 以顾客为关注焦点\n"
                "   - 质量方针制定\n"
                "   - 策划（质量目标、体系策划）\n"
                "   - 职责、权限和沟通\n"
                "   - 管理评审\n\n"
                "6. 资源管理\n"
                "   - 人力资源（培训、能力、意识）\n"
                "   - 基础设施（厂房、设备、软件）\n"
                "   - 工作环境（洁净度、温湿度、静电防护）\n\n"
                "7. 产品实现\n"
                "   - 产品实现的策划\n"
                "   - 与顾客有关的过程\n"
                "   - 设计和开发（结合有源医疗器械特点）\n"
                "   - 采购（关键元器件、供应商管理）\n"
                "   - 生产和服务提供（含软件开发）\n"
                "   - 监视和测量设备的控制\n\n"
                "8. 测量、分析和改进\n"
                "   - 监视和测量（顾客满意、内部审核、过程监测）\n"
                "   - 不合格品控制\n"
                "   - 数据分析\n"
                "   - 改进（纠正措施、预防措施）\n\n"
                "9. 风险管理\n"
                "   - 风险管理在质量管理体系中的应用\n"
                "   - 产品生命周期各阶段的风险管理活动\n\n"
                "10. 文件控制\n"
                "    - 文件编写、审批、发放、更改、作废流程\n\n"
                "11. 记录控制\n"
                "    - 记录的标识、存储、保护、检索、保留、处置\n\n"
                "【格式要求】\n"
                "- 使用正式的质量手册格式，每个章节要有详细内容\n"
                "- 包含ISO 13485和中国法规的双重条款号引用\n"
                "- 使用Markdown格式，包含标题层级、列表、表格\n"
                "- 文件编号格式：QM-001（质量手册）\n"
                "- 每个过程要描述：输入、输出、职责、流程、相关文件\n\n"
                "【占位符规范】\n"
                "使用 [需填写：xxx] 标记需要用户填写的内容\n"
            )
            
            result = await engine.generate_document(
                system_prompt=system_prompt,
                user_message=f"请生成质量手册。{product_context}{company_context}",
                max_iterations=80,
            )
            # Remove any introductory text before the actual content
            result = _clean_generated_content(result)
            state.generated_docs["qms_01_quality_manual"] = result
            logger.info("Generated QMS Quality Manual")
        
        # 2. Generate ALL Procedure Documents - DETAILED VERSION
        if chapters.get("procedures"):
            for proc in chapters["procedures"]["documents"]:
                system_prompt = (
                    f"你是一名资深的医疗器械质量管理体系专家。\n"
                    f"请生成'{proc['title']}'程序文件。\n\n"
                    "【重要要求】\n"
                    "- 直接输出程序文件正文，不要任何前言、说明\n"
                    "- 内容要详细、专业、可直接用于审核\n"
                    "- 结合公司实际产品特点\n"
                    "- 每个步骤都要详细描述操作方法\n\n"
                    f"【条款号】ISO 13485:2016 {proc['clause']}\n"
                    f"【目的】{proc['purpose']}\n"
                    f"{regulations_ref}\n\n"
                    f"【程序文件必须包含的章节】\n"
                    f"1. 目的（详细说明本程序要达到的目的）\n"
                    f"2. 范围（适用于哪些产品、哪些部门、哪些过程）\n"
                    f"3. 引用文件（相关法规、标准、内部文件）\n"
                    f"4. 职责（详细说明各部门/岗位的职责权限）\n"
                    f"5. 程序内容（详细的操作步骤，每步都要描述清楚）\n"
                    f"6. 相关文件（本程序引用和产生的文件清单）\n"
                    f"7. 记录（本程序产生的所有记录表格清单）\n"
                    f"8. 附表（空白表格模板）\n\n"
                    f"【格式要求】\n"
                    f"- 文件编号格式：XX-QP-XXX\n"
                    f"- 包含版本号、修改记录\n"
                    f"- 步骤要详细到可以直接执行\n"
                    f"- 使用Markdown格式\n\n"
                    f"【占位符规范】\n"
                    f"使用 [需填写：xxx] 标记需要用户填写的内容\n"
                )
                
                result = await engine.generate_document(
                    system_prompt=system_prompt,
                    user_message=f"请生成'{proc['title']}'程序文件。{product_context}{company_context}",
                    max_iterations=60,
                )
                result = _clean_generated_content(result)
                state.generated_docs[f"qms_02_{proc['code']}"] = result
                logger.info(f"Generated QMS Procedure: {proc['code']}")
        
        # 3. Generate Work Instructions - DETAILED VERSION
        if chapters.get("work_instructions"):
            for category in chapters["work_instructions"]["categories"]:
                system_prompt = (
                    f"你是一名资深的医疗器械质量管理体系专家。\n"
                    f"请生成'{category['category']}'类别的作业指导书。\n\n"
                    "【重要要求】\n"
                    "- 直接输出作业指导书正文，不要任何前言、说明\n"
                    "- 内容要详细到操作人员可以直接按步骤执行\n"
                    "- 结合公司实际产品特点\n\n"
                    f"【包含的作业指导书】\n"
                )
                for doc in category["documents"]:
                    system_prompt += f"- {doc}\n"
                
                system_prompt += (
                    f"\n【每份作业指导书必须包含】\n"
                    f"1. 文件编号和版本号\n"
                    f"2. 文件名称\n"
                    f"3. 目的\n"
                    f"4. 适用范围（产品、工序）\n"
                    f"5. 所需设备/工具/材料\n"
                    f"6. 详细操作步骤（图文说明位置）\n"
                    f"7. 质量控制要点\n"
                    f"8. 注意事项和安全要求\n"
                    f"9. 相关记录\n\n"
                    f"【格式要求】\n"
                    f"- 文件编号格式：XX-WI-XXX\n"
                    f"- 步骤要编号，详细到可以直接执行\n"
                    f"- 使用Markdown格式\n\n"
                    f"【占位符规范】\n"
                    f"使用 [需填写：xxx] 标记需要用户填写的内容\n"
                )
                
                result = await engine.generate_document(
                    system_prompt=system_prompt,
                    user_message=f"请生成'{category['category']}'类别的作业指导书。{product_context}{company_context}",
                    max_iterations=60,
                )
                result = _clean_generated_content(result)
                state.generated_docs[f"qms_03_{category['category']}"] = result
                logger.info(f"Generated QMS Work Instructions: {category['category']}")
        
        # 4. Generate Quality Records Templates
        if chapters.get("quality_records"):
            system_prompt = (
                "你是一名资深的医疗器械质量管理体系专家。\n"
                "请生成质量记录表格模板。\n\n"
                "【重要要求】\n"
                "- 直接输出表格模板，不要任何说明\n"
                "- 表格要专业、完整、可直接使用\n\n"
                "【包含的记录表格】\n"
            )
            for form in chapters["quality_records"]["forms"]:
                system_prompt += f"- {form}\n"
            
            system_prompt += (
                "\n【每份表格必须包含】\n"
                "1. 表格标题\n"
                "2. 表格编号（XX-FM-XXX格式）\n"
                "3. 版本号\n"
                "4. 完整的表格结构（表头、列、行）\n"
                "5. 填写说明\n"
                "6. 审批签名栏\n\n"
                "【格式要求】\n"
                "使用Markdown表格格式\n"
            )
            
            result = await engine.generate_document(
                system_prompt=system_prompt,
                user_message=f"请生成质量记录表格模板。{product_context}{company_context}",
                max_iterations=50,
            )
            result = _clean_generated_content(result)
            state.generated_docs["qms_04_quality_records"] = result
            logger.info("Generated QMS Quality Records")
        
        logger.info("QMS document generation complete")
    
    bg.add_task(_generate_qms_docs, )
    bg.add_task(_generate_qms_docs, )
    
    return {
        "task_id": str(uuid.uuid4()),
        "status": "started",
        "message": "QMS document generation started",
    }


@router.get("/qms-tree")
async def get_qms_tree():
    """Get QMS document tree structure."""
    import json as json_mod
    from pathlib import Path
    
    qms_file = Path(__file__).parent.parent.parent.parent / "regulations" / "iso13485.json"
    qms_requirements = {}
    if qms_file.exists():
        with open(qms_file, 'r', encoding='utf-8') as f:
            qms_requirements = json_mod.load(f)
    
    chapters = qms_requirements.get("iso13485", {}).get("chapters", {})
    
    tree = {
        "title": "ISO 13485 质量管理体系文件",
        "children": []
    }
    
    # Quality Manual
    if chapters.get("quality_manual"):
        tree["children"].append({
            "id": "qms_quality_manual",
            "title": "一级文件：质量手册",
            "required": True,
        })
    
    # Procedures
    if chapters.get("procedures"):
        proc_node = {
            "id": "qms_procedures",
            "title": "二级文件：程序文件",
            "required": True,
            "children": []
        }
        for proc in chapters["procedures"]["documents"]:
            proc_node["children"].append({
                "id": f"qms_proc_{proc['code']}",
                "title": f"{proc['code']} {proc['title']}",
                "required": True,
            })
        tree["children"].append(proc_node)
    
    # Work Instructions
    if chapters.get("work_instructions"):
        wi_node = {
            "id": "qms_work_instructions",
            "title": "三级文件：作业指导书",
            "required": True,
            "children": []
        }
        for category in chapters["work_instructions"]["categories"]:
            wi_node["children"].append({
                "id": f"qms_wi_{category['category']}",
                "title": category["category"],
                "required": True,
            })
        tree["children"].append(wi_node)
    
    return tree


@router.post("/generate", response_model=DocumentGenerateResponse)
async def generate_document(req: DocumentGenerateRequest, bg: BackgroundTasks):
    """Trigger document generation."""
    state = get_state()

    if state.workspace_path is None:
        raise HTTPException(status_code=400, detail="Workspace not initialized. Call /workspace/init first.")

    task_id = str(uuid.uuid4())

    bg.add_task(_generate_node_bg, req.node_id, task_id)

    return DocumentGenerateResponse(
        task_id=task_id,
        status="started",
        message=f"Generation started for node '{req.node_id}' in '{req.mode}' mode",
    )


@router.get("/export/{node_id}")
async def export_document(node_id: str, format: str = "md"):
    """Export a single document as Word (docx) with proper formatting."""
    from fastapi.responses import StreamingResponse
    import io
    import re
    
    state = get_state()
    
    if node_id not in state.generated_docs:
        raise HTTPException(status_code=404, detail=f"Document '{node_id}' not found or not generated yet.")
    
    content = state.generated_docs[node_id]
    
    if format == "docx":
        try:
            from docx import Document
            from docx.shared import Pt, Inches, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
            
            doc = Document()
            
            # Set default font for Chinese
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Microsoft YaHei'
            font.size = Pt(11)
            
            # Set page margins
            for section in doc.sections:
                section.top_margin = Cm(2.54)
                section.bottom_margin = Cm(2.54)
                section.left_margin = Cm(3.18)
                section.right_margin = Cm(3.18)
            
            # Process content
            lines = content.split('\n')
            i = 0
            while i < len(lines):
                line = lines[i]
                stripped = line.strip()
                
                # Skip empty lines
                if not stripped:
                    i += 1
                    continue
                
                # Handle headings
                if stripped.startswith('# '):
                    p = doc.add_heading(stripped[2:], level=1)
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif stripped.startswith('## '):
                    p = doc.add_heading(stripped[3:], level=2)
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif stripped.startswith('### '):
                    p = doc.add_heading(stripped[4:], level=3)
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif stripped.startswith('#### '):
                    p = doc.add_heading(stripped[5:], level=4)
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Handle tables
                elif stripped.startswith('|') and '|' in stripped[1:]:
                    # Collect all table rows
                    table_lines = []
                    while i < len(lines) and lines[i].strip().startswith('|'):
                        row_line = lines[i].strip()
                        # Skip separator lines (|---|---|)
                        if not re.match(r'^\|[\s\-:]+\|', row_line):
                            table_lines.append(row_line)
                        i += 1
                    i -= 1  # Adjust for outer loop increment
                    
                    if table_lines:
                        # Parse table
                        rows = []
                        for tl in table_lines:
                            cells = [c.strip() for c in tl.split('|')[1:-1]]
                            rows.append(cells)
                        
                        if rows:
                            # Create Word table
                            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                            table.style = 'Table Grid'
                            table.alignment = WD_TABLE_ALIGNMENT.CENTER
                            
                            for r_idx, row in enumerate(rows):
                                for c_idx, cell_text in enumerate(row):
                                    if c_idx < len(table.rows[r_idx].cells):
                                        cell = table.rows[r_idx].cells[c_idx]
                                        # Remove markdown bold
                                        clean_text = cell_text.replace('**', '')
                                        cell.text = clean_text
                                        # Style header row
                                        if r_idx == 0:
                                            for paragraph in cell.paragraphs:
                                                for run in paragraph.runs:
                                                    run.bold = True
                    
                    doc.add_paragraph()  # Add space after table
                
                # Handle bullet lists
                elif stripped.startswith('- ') or stripped.startswith('* '):
                    text = stripped[2:]
                    # Remove markdown formatting
                    text = text.replace('**', '').replace('*', '')
                    p = doc.add_paragraph(text, style='List Bullet')
                
                # Handle numbered lists
                elif re.match(r'^\d+\.\s', stripped):
                    text = re.sub(r'^\d+\.\s', '', stripped)
                    text = text.replace('**', '').replace('*', '')
                    p = doc.add_paragraph(text, style='List Number')
                
                # Handle blockquotes
                elif stripped.startswith('>'):
                    text = stripped.lstrip('> ')
                    text = text.replace('**', '')
                    p = doc.add_paragraph(text)
                    p.paragraph_format.left_indent = Cm(1)
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(100, 100, 100)
                
                # Handle horizontal rules
                elif stripped in ('---', '***', '___'):
                    p = doc.add_paragraph()
                    p.add_run('─' * 50)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Regular paragraph
                else:
                    # Remove markdown formatting
                    clean_text = stripped.replace('**', '').replace('*', '').replace('`', '')
                    
                    p = doc.add_paragraph()
                    # Handle inline formatting
                    parts = re.split(r'(\*\*.*?\*\*)', clean_text)
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            run = p.add_run(part[2:-2])
                            run.bold = True
                        else:
                            p.add_run(part)
                
                i += 1
            
            # Save to bytes
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            # Generate proper filename
            import re as re_mod
            clean_name = re_mod.sub(r'[^\w\s\-]', '', node_id).strip()
            clean_name = re_mod.sub(r'\s+', '_', clean_name)
            filename = f"{clean_name}.docx"
            
            return StreamingResponse(
                buffer,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f'attachment; filename="{filename}"'}
            )
        except ImportError:
            raise HTTPException(status_code=500, detail="python-docx not installed")
    else:
        # Return Markdown
        filename = f"{node_id}.md"
        return StreamingResponse(
            io.BytesIO(content.encode('utf-8')),
            media_type="text/markdown",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )


@router.get("/export-all")
async def export_all_documents(format: str = "docx"):
    """Export all generated documents as a zip file with proper Word formatting."""
    from fastapi.responses import StreamingResponse
    import io
    import zipfile
    import re
    
    state = get_state()
    
    if not state.generated_docs:
        raise HTTPException(status_code=404, detail="No documents generated yet.")
    
    buffer = io.BytesIO()
    
    with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
        for node_id, content in state.generated_docs.items():
            if format == "docx":
                try:
                    from docx import Document
                    from docx.shared import Pt
                    
                    doc = Document()
                    style = doc.styles['Normal']
                    font = style.font
                    font.name = 'Microsoft YaHei'
                    font.size = Pt(11)
                    
                    lines = content.split('\n')
                    for line in lines:
                        if line.startswith('# '):
                            doc.add_heading(line[2:], level=1)
                        elif line.startswith('## '):
                            doc.add_heading(line[3:], level=2)
                        elif line.startswith('### '):
                            doc.add_heading(line[4:], level=3)
                        elif line.startswith('- ') or line.startswith('* '):
                            doc.add_paragraph(line[2:], style='List Bullet')
                        elif line.strip():
                            p = doc.add_paragraph()
                            parts = line.split('**')
                            for i, part in enumerate(parts):
                                if i % 2 == 1:
                                    run = p.add_run(part)
                                    run.bold = True
                                else:
                                    p.add_run(part)
                    
                    doc_buffer = io.BytesIO()
                    doc.save(doc_buffer)
                    zf.writestr(f"{node_id}.docx", doc_buffer.getvalue())
                except ImportError:
                    # Fallback to markdown
                    zf.writestr(f"{node_id}.md", content)
            else:
                zf.writestr(f"{node_id}.md", content)
    
    buffer.seek(0)
    
    template_name = state.config.fda_template
    return StreamingResponse(
        buffer,
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{template_name}_documents.zip"'}
    )


@router.get("/list")
async def list_documents():
    """List all generated documents with their status."""
    state = get_state()
    
    docs = []
    for node_id, content in state.generated_docs.items():
        docs.append({
            "node_id": node_id,
            "content_length": len(content),
            "preview": content[:200] + "..." if len(content) > 200 else content,
        })
    
    return {
        "total": len(docs),
        "documents": docs,
    }
